from __future__ import annotations

import re
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from pypdf import PdfReader


SOURCE_PDF = Path(r"C:\AIroom\codexcli\財務諸表理論アプリ\assets\pdfs\財務諸表論問題集.pdf")
OUT_DIR = Path(r"C:\Users\kawat\Documents\Codex\2026-05-29\quiz_database")
DOCX_PATH = OUT_DIR / "財務諸表論問題集_第1回_game_trial.docx"
PDF_PATH = OUT_DIR / "財務諸表論問題集_第1回_game_trial.pdf"
XLSX_PATH = OUT_DIR / "財務諸表論問題集_第1回_pdf_game.xlsx"


def normalize_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def clean_explanation(text: str) -> str:
    text = normalize_text(text)
    return re.sub(r"\s+[45]$", "", text).strip()


def parse_round1_items() -> list[dict[str, str]]:
    reader = PdfReader(str(SOURCE_PDF))
    text = " ".join((reader.pages[i - 1].extract_text() or "").replace("\n", " ") for i in [4, 5])
    text = normalize_text(text)
    chunks = []
    pattern = re.compile(r"(?:^|\s)(\d{1,2})\s+(.+?)(?=\s+\d{1,2}\s+|$)")
    for match in pattern.finditer(text):
        number = int(match.group(1))
        if not 1 <= number <= 12:
            continue
        chunk = match.group(2).strip()
        split = re.search(r"\s([\u25cb\u00d7])(?:\s|$)", chunk)
        if not split:
            continue
        problem = normalize_text(chunk[: split.start()])
        answer = split.group(1)
        explanation = clean_explanation(chunk[split.end() :])
        if problem and answer:
            chunks.append({
                "id": f"基本01-{number:03d}",
                "number": str(number),
                "title": "第1回 財務会計の意義と機能",
                "problem": problem,
                "answer": answer,
                "explanation": explanation,
            })
    chunks = sorted({item["id"]: item for item in chunks}.values(), key=lambda item: int(item["number"]))
    if len(chunks) != 12:
        raise RuntimeError(f"expected 12 items, got {len(chunks)}")
    return chunks


def set_japanese_font(run, name: str = "Yu Gothic") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_page_number(paragraph, current: int) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(str(current))
    set_japanese_font(run)
    run.font.size = Pt(9)


def set_paragraph_text(paragraph, text: str, size: float = 12, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    set_japanese_font(run)
    run.font.size = Pt(size)
    run.bold = bold


def add_labeled_block(doc: Document, label: str, text: str) -> None:
    label_p = doc.add_paragraph()
    set_paragraph_text(label_p, label, 10, True)
    body = doc.add_paragraph()
    body.paragraph_format.line_spacing = 1.25
    set_paragraph_text(body, text or "（解説なし）", 12)


def add_page_break(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_docx(items: list[dict[str, str]]) -> None:
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Yu Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    styles["Normal"].font.size = Pt(11)

    total_problem_pages = len(items)
    page_no = 1
    for item in items:
        title = doc.add_paragraph()
        set_paragraph_text(title, f"{item['id']}　{item['title']}", 13, True)
        subtitle = doc.add_paragraph()
        set_paragraph_text(subtitle, "問題", 11, True)
        add_labeled_block(doc, "問題文", item["problem"])
        hint = doc.add_paragraph()
        set_paragraph_text(hint, "解答欄：○ または × を入力してください。", 10)
        add_page_number(doc.add_paragraph(), page_no)
        page_no += 1
        add_page_break(doc)

    for item in items:
        title = doc.add_paragraph()
        set_paragraph_text(title, f"{item['id']}　{item['title']}", 13, True)
        subtitle = doc.add_paragraph()
        set_paragraph_text(subtitle, "解答・解説", 11, True)
        add_labeled_block(doc, "問題文", item["problem"])
        add_labeled_block(doc, "正解", item["answer"])
        add_labeled_block(doc, "解説", item["explanation"])
        add_page_number(doc.add_paragraph(), page_no)
        page_no += 1
        if item != items[-1]:
            add_page_break(doc)

    doc.save(DOCX_PATH)


def build_excel(items: list[dict[str, str]]) -> None:
    wb = openpyxl.Workbook()
    ws_answers = wb.active
    ws_answers.title = "問題"
    ws_pages = wb.create_sheet("ページ情報")
    ws_usage = wb.create_sheet("使い方")

    answer_headers = ["問題番号"]
    for i in range(1, 21):
        answer_headers += [f"{i}問目の回答範囲", f"{i}問目の正解"]
    page_headers = [
        "問題番号",
        "PDF内問題開始ページ",
        "PDF内問題終了ページ",
        "PDF内解答開始ページ",
        "PDF内解答終了ページ",
        "書籍印刷問題開始ページ",
        "書籍印刷問題終了ページ",
        "書籍印刷解答開始ページ",
        "書籍印刷解答終了ページ",
    ]
    ws_answers.append(answer_headers)
    ws_pages.append(page_headers)

    problem_count = len(items)
    for idx, item in enumerate(items, start=1):
        ws_answers.append([item["id"], "正誤", f"{item['answer']}|{'まる' if item['answer'] == '○' else 'ばつ'}|{'o' if item['answer'] == '○' else 'x'}"] + [None] * 38)
        answer_page = problem_count + idx
        ws_pages.append([item["id"], idx, idx, answer_page, answer_page, idx, idx, answer_page, answer_page])

    ws_usage.append(["財務諸表論問題集 第1回 Word/PDF試作用データ", "PDF + Excel モード"])
    ws_usage.append(["対象PDF", PDF_PATH.name])
    ws_usage.append(["対象範囲", "第1回 財務会計の意義と機能"])
    ws_usage.append(["問題数", len(items)])
    ws_usage.append(["構成", "PDF前半12ページが問題、後半12ページが解答・解説。"])

    style_workbook(wb)
    wb.save(XLSX_PATH)


def style_workbook(wb: openpyxl.Workbook) -> None:
    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(style="thin", color="D9E2F3")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for ws in wb.worksheets:
        ws.freeze_panes = "A2"
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in ws.iter_rows():
            for cell in row:
                cell.border = border
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        for col in range(1, ws.max_column + 1):
            width = 14
            if col == 1:
                width = 18
            elif ws.title == "問題":
                width = 18 if col % 2 == 0 else 16
            elif ws.title == "ページ情報":
                width = 20
            ws.column_dimensions[get_column_letter(col)].width = width
        ws.auto_filter.ref = ws.dimensions


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    items = parse_round1_items()
    build_docx(items)
    build_excel(items)
    print(DOCX_PATH)
    print(XLSX_PATH)
    print(f"items={len(items)}")
    for item in items[:3]:
        print(item["id"], item["answer"], item["problem"][:50])


if __name__ == "__main__":
    main()
